import sys
try:
    from docx import Document
except ImportError:
    print("python-docx is not installed")
    sys.exit(1)

doc = Document(r"C:\Users\njoki\.gemini\antigravity\scratch\laptop-ecommerce-frontend\data\NEXGEN SOLUTIONS PRICING LIST  1 u.docx")
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text.strip())

for table in doc.tables:
    for row in table.rows:
        print([cell.text.strip() for cell in row.cells])
